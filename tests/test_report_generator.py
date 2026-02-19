import os
import pytest
from datetime import date
from docx import Document
from src.services.report_generator import ReportGenerator
from src.services.project_service import get_circled_number
from src.db_schema.orm_models import (
    Project,
    CabinetTemplate,
    CabinetPart,
    ProjectCabinet,
    ProjectCabinetAccessory,
    ProjectCabinetPart,
    Accessory,
)


@pytest.fixture
def sample_project_orm():
    """
    Given: a SQLAlchemy Project with one cabinet and one accessory
    """
    # Project setup
    project = Project(
        name="ORM Project",
        kitchen_type="LOFT",
        order_number="ORM001",
        client_name="ORM Client",
        client_address="456 ORM Ave",
        client_phone="555-ORM",
        client_email="orm@example.com",
        blaty=True,
        blaty_note="Blaty test",
        cokoly=True,
        cokoly_note="Cokoły test",
        uwagi=True,
        uwagi_note="Uwagi test",
    )
    # CabinetTemplate with basic info
    ct = CabinetTemplate(
        kitchen_type="LOFT",
        name="Standard",
    )

    # Add cabinet parts with different materials to trigger all report sections
    parts = [
        CabinetPart(
            part_name="bok lewy",
            height_mm=720,
            width_mm=500,
            pieces=1,
            material="PLYTA 18",
            wrapping="D",
        ),
        CabinetPart(
            part_name="bok prawy",
            height_mm=720,
            width_mm=500,
            pieces=1,
            material="PLYTA 18",
            wrapping="D",
        ),
        CabinetPart(
            part_name="polka",
            height_mm=480,
            width_mm=500,
            pieces=2,
            material="PLYTA 18",
            wrapping="D",
        ),
        CabinetPart(
            part_name="front",
            height_mm=700,
            width_mm=596,
            pieces=1,
            material="FRONT 16",
            wrapping="DDKK",
        ),
        CabinetPart(
            part_name="plecy",
            height_mm=715,
            width_mm=580,
            pieces=1,
            material="HDF 3",
        ),
    ]
    ct.parts = parts

    # One cabinet of quantity 2
    cab = ProjectCabinet(
        sequence_number=1,
        body_color="Oak",
        front_color="Maple",
        handle_type="KROMA",
        quantity=2,
    )
    # assign relationships
    cab.project = project
    cab.cabinet_type = ct
    project.cabinets = [cab]

    # Accessory and link
    acc = Accessory(name="Hinge X")
    link = ProjectCabinetAccessory(count=4)
    # assign relationships
    link.project_cabinet = cab
    link.accessory = acc
    cab.accessories = [link]

    return project


def test_generate_creates_file(tmp_path, sample_project_orm):
    """
    Given: a ReportGenerator and ORM project
    When: generate is called
    Then: output file exists with correct name
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    assert os.path.isfile(output)
    assert output.endswith("projekt_ORM001.docx")


def test_generate_uses_incremented_filename_when_file_exists(
    tmp_path, sample_project_orm
):
    """
    Given: a target filename that already exists
    When: generate is called
    Then: report is saved using incremented suffix "(1)"
    """
    existing = tmp_path / "projekt_ORM001.docx"
    existing.write_bytes(b"existing")

    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)

    assert os.path.isfile(output)
    assert output.endswith("projekt_ORM001 (1).docx")


def test_generate_sanitizes_invalid_characters_in_filename(tmp_path):
    """
    Given: order number with filesystem-invalid characters
    When: generate is called
    Then: output filename is sanitized and report is saved
    """
    project = Project(
        name="Sanitize Project",
        kitchen_type="LOFT",
        order_number="GAMA 9080/13",
        client_name="Client",
        client_address="Address",
        client_phone="555-000",
        client_email="client@example.com",
    )
    project.cabinets = []

    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)

    assert os.path.isfile(output)
    assert output.endswith("projekt_GAMA 9080_13.docx")


def test_header_contains_orm_metadata(tmp_path, sample_project_orm):
    """
    Given: a generated report
    When: reading the header table
    Then: it contains ORM project metadata and today's date
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)
    header = doc.sections[0].header
    text = header.tables[0].cell(0, 0).text
    assert "Klient: ORM Client" in text  # Polish localization
    assert "Nr zamówienia: ORM001" in text  # Polish localization
    assert date.today().isoformat() in text


def test_body_tables_count_and_headers(tmp_path, sample_project_orm):
    """
    Given: a report with derived parts
    When: collecting body tables
    Then: there are four tables titled with color-specific headings for FORMATKI/FRONTY
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    # ---- Given / When: generated the document ----

    # Then: there should be exactly four Heading 2 titles in this order
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert headings == [
        "FORMATKI (PLYTA 18) - Oak",
        "FRONTY - Maple",
        "HDF",
        "AKCESORIA",
    ]

    # And: the four body tables have the correct column headers
    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    expected_hdr = [
        ["Lp.", "Nazwa", "Wymiary (mm)", "Ilość", "Okleina", "Kolor", "Uwagi"],
        ["Lp.", "Nazwa", "Wymiary (mm)", "Ilość", "Okleina", "Kolor", "Uwagi"],
        ["Lp.", "Nazwa", "Wymiary (mm)", "Ilość", "Okleina", "Kolor", "Uwagi"],
        ["Poz.", "Nazwa akcesorium", "Ilość", "Uwagi"],
    ]
    for table, hdr in zip(body_tables, expected_hdr):
        assert [cell.text for cell in table.rows[0].cells] == hdr


def test_derived_formatki_quantities(tmp_path, sample_project_orm):
    """
    Given: derived FORMATKI table
    When: reading its rows
    Then: the quantities reflect count * cabinet.quantity
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]
    fmt_table = body_tables[0]

    # Header + 3 data rows (bok, wieniec, polka)
    assert len(fmt_table.rows) == 4

    # Quantity is now in cell index 3 (after "Lp.", "Nazwa" and "Wymiary")
    quantities = [int(row.cells[3].text) for row in fmt_table.rows[1:]]
    assert quantities == [2, 2, 4]  # (1*2), (1*2), (2*2)


def test_plyta_16_hides_color_values_but_keeps_column(tmp_path):
    """
    Given: formatki in both PLYTA 16 and PLYTA 18 sections
    When: generating report
    Then: PLYTA 16 rows have empty "Kolor", while PLYTA 18 still shows color
    """
    project = Project(
        name="PLYTA 16 Color Test",
        kitchen_type="LOFT",
        order_number="PLYTA16-001",
        client_name="Client",
        client_address="Address",
        client_phone="555-000",
        client_email="client@example.com",
    )

    ct = CabinetTemplate(kitchen_type="LOFT", name="Panel Template")
    ct.parts = [
        CabinetPart(
            part_name="Panel 16",
            height_mm=720,
            width_mm=500,
            pieces=1,
            material="PLYTA 16",
            wrapping="D",
        ),
        CabinetPart(
            part_name="Panel 18",
            height_mm=720,
            width_mm=500,
            pieces=1,
            material="PLYTA 18",
            wrapping="D",
        ),
    ]

    cab = ProjectCabinet(
        sequence_number=1,
        body_color="White",
        front_color="Oak",
        handle_type="KROMA",
        quantity=1,
    )
    cab.project = project
    cab.cabinet_type = ct
    project.cabinets = [cab]

    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    plyta_16_table = None
    plyta_18_table = None

    for table in body_tables:
        part_names = [row.cells[1].text.strip() for row in table.rows[1:]]
        if "Panel 16" in part_names:
            plyta_16_table = table
        if "Panel 18" in part_names:
            plyta_18_table = table

    assert plyta_16_table is not None
    assert plyta_18_table is not None

    # Column still exists
    assert plyta_16_table.rows[0].cells[5].text == "Kolor"

    # PLYTA 16 hides values
    plyta_16_row = next(
        row
        for row in plyta_16_table.rows[1:]
        if row.cells[1].text.strip() == "Panel 16"
    )
    assert plyta_16_row.cells[5].text.strip() == ""

    # Other sections keep color values
    plyta_18_row = next(
        row
        for row in plyta_18_table.rows[1:]
        if row.cells[1].text.strip() == "Panel 18"
    )
    assert plyta_18_row.cells[5].text.strip() == "White"


def test_color_grouped_sections_for_selected_sections(tmp_path):
    """
    Given: FRONTY, PLYTA 12 and PLYTA 18 with multiple colors
    When: generating report
    Then: each color is rendered in a separate section and rows are sorted by LP in each group
    """
    project = Project(
        name="Color Grouping Project",
        kitchen_type="LOFT",
        order_number="COLOR-GROUP-001",
        client_name="Client",
        client_address="Address",
        client_phone="555-000",
        client_email="client@example.com",
    )

    ct = CabinetTemplate(kitchen_type="LOFT", name="Color Group Template")
    ct.parts = [
        CabinetPart(
            part_name="Panel 18",
            height_mm=720,
            width_mm=500,
            pieces=1,
            material="PLYTA 18",
            wrapping="D",
        ),
        CabinetPart(
            part_name="Panel 12",
            height_mm=300,
            width_mm=500,
            pieces=1,
            material="PLYTA 12",
            wrapping="D",
        ),
        CabinetPart(
            part_name="Front drzwi",
            height_mm=700,
            width_mm=500,
            pieces=1,
            material="FRONT 18",
            wrapping="DD",
        ),
    ]

    cab_1 = ProjectCabinet(
        sequence_number=1,
        body_color="White",
        front_color="Maple",
        handle_type="A",
        quantity=1,
    )
    cab_2 = ProjectCabinet(
        sequence_number=2,
        body_color="Gray",
        front_color="Black",
        handle_type="B",
        quantity=1,
    )
    cab_3 = ProjectCabinet(
        sequence_number=3,
        body_color="White",
        front_color="Maple",
        handle_type="C",
        quantity=1,
    )
    for cab in (cab_1, cab_2, cab_3):
        cab.project = project
        cab.cabinet_type = ct
    project.cabinets = [cab_1, cab_2, cab_3]

    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert "FORMATKI (PLYTA 18) - Gray" in headings
    assert "FORMATKI (PLYTA 18) - White" in headings
    assert "FORMATKI (PLYTA 12) - Gray" in headings
    assert "FORMATKI (PLYTA 12) - White" in headings
    assert "FRONTY - Black" in headings
    assert "FRONTY - Maple" in headings

    def find_tables_with_name(name: str):
        return [
            table
            for table in body_tables
            if any(row.cells[1].text.strip() == name for row in table.rows[1:])
        ]

    plyta_18_tables = find_tables_with_name("Panel 18")
    plyta_12_tables = find_tables_with_name("Panel 12")
    fronty_tables = find_tables_with_name("Front drzwi")

    assert len(plyta_18_tables) == 2
    assert len(plyta_12_tables) == 2
    assert len(fronty_tables) == 2

    seq1 = get_circled_number(1)
    seq2 = get_circled_number(2)
    seq3 = get_circled_number(3)

    # PLYTA 18: Gray group has seq2; White group has seq1 then seq3
    gray_plyta_18 = next(
        t for t in plyta_18_tables if t.rows[1].cells[5].text.strip() == "Gray"
    )
    white_plyta_18 = next(
        t for t in plyta_18_tables if t.rows[1].cells[5].text.strip() == "White"
    )
    assert [r.cells[0].text.strip() for r in gray_plyta_18.rows[1:]] == [seq2]
    assert [r.cells[0].text.strip() for r in white_plyta_18.rows[1:]] == [seq1, seq3]

    # PLYTA 12: Gray group has seq2; White group has seq1 then seq3
    gray_plyta_12 = next(
        t for t in plyta_12_tables if t.rows[1].cells[5].text.strip() == "Gray"
    )
    white_plyta_12 = next(
        t for t in plyta_12_tables if t.rows[1].cells[5].text.strip() == "White"
    )
    assert [r.cells[0].text.strip() for r in gray_plyta_12.rows[1:]] == [seq2]
    assert [r.cells[0].text.strip() for r in white_plyta_12.rows[1:]] == [seq1, seq3]

    # FRONTY: Black group has seq2; Maple group has seq1 then seq3
    black_fronty = next(
        t for t in fronty_tables if t.rows[1].cells[5].text.strip() == "Black"
    )
    maple_fronty = next(
        t for t in fronty_tables if t.rows[1].cells[5].text.strip() == "Maple"
    )
    assert [r.cells[0].text.strip() for r in black_fronty.rows[1:]] == [seq2]
    assert [r.cells[0].text.strip() for r in maple_fronty.rows[1:]] == [seq1, seq3]


def test_accessories_section(tmp_path, sample_project_orm):
    """
    Given: derived AKCESORIA table
    When: reading its single row
    Then: it matches accessory name and count
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)
    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]
    acc_table = body_tables[3]
    assert len(acc_table.rows) == 2
    cells = acc_table.rows[1].cells
    # Name at index 1, quantity at 2, notes at 3
    assert cells[1].text == "Hinge X"
    assert cells[2].text == str(4 * sample_project_orm.cabinets[0].quantity)
    assert cells[3].text == ""


def test_should_break_page_when_section_would_not_fit():
    """
    Given: a page almost full
    When: a section that normally fits one page is about to be added
    Then: a page break is required to keep the section together
    """
    rg = ReportGenerator()
    doc = Document()

    # Leave very little estimated space on current page.
    for _ in range(43):
        doc.add_paragraph("x")

    assert rg._should_break_page_for_section(doc, items_count=5) is True


def test_should_not_break_page_when_section_fits():
    """
    Given: enough estimated space left on the page
    When: adding a small section
    Then: no page break is needed
    """
    rg = ReportGenerator()
    doc = Document()

    for _ in range(10):
        doc.add_paragraph("x")

    assert rg._should_break_page_for_section(doc, items_count=5) is False


def test_should_break_page_respects_strictness_setting():
    """
    Given: the same amount of free space on page
    When: changing strictness setting
    Then: break decision changes with strictness level
    """

    class _FakeSettingsService:
        def __init__(self, value: str):
            self.value = value

        def get_setting_value(self, key: str, default=None):
            if key == "report_page_break_strictness":
                return self.value
            return default

    doc = Document()
    for _ in range(20):
        doc.add_paragraph("x")

    rg = ReportGenerator()

    rg.settings_service = _FakeSettingsService("Lagodna")
    assert rg._should_break_page_for_section(doc, items_count=5) is False

    rg.settings_service = _FakeSettingsService("Standardowa")
    assert rg._should_break_page_for_section(doc, items_count=5) is False

    rg.settings_service = _FakeSettingsService("Ostra")
    assert rg._should_break_page_for_section(doc, items_count=5) is True


def test_constructor_accepts_legacy_positional_session(session):
    """
    Given: legacy constructor call with Session as first positional argument
    When: creating ReportGenerator
    Then: db-backed services are correctly initialized
    """
    rg = ReportGenerator(session)
    assert rg.project_service is not None
    assert rg.settings_service is not None


def test_report_contains_glass_shelves_section(tmp_path):
    """
    Given: a cabinet that contains a glass shelf part
    When: generating the report
    Then: the report has a dedicated "PÓŁKI SZKLANE" section and row
    """
    project = Project(
        name="Glass Shelf Project",
        kitchen_type="LOFT",
        order_number="GLASS-001",
        client_name="Glass Client",
        client_address="Glass Street 1",
        client_phone="555-GLASS",
        client_email="glass@example.com",
    )

    ct = CabinetTemplate(kitchen_type="LOFT", name="Glass Template")
    ct.parts = [
        CabinetPart(
            part_name="półka szklana",
            height_mm=480,
            width_mm=300,
            pieces=1,
            material="PÓŁKA SZKLANA",
        ),
        CabinetPart(
            part_name="front",
            height_mm=700,
            width_mm=596,
            pieces=1,
            material="FRONT 16",
        ),
    ]

    cab = ProjectCabinet(
        sequence_number=1,
        body_color="Oak",
        front_color="Maple",
        handle_type="KROMA",
        quantity=2,
    )
    cab.project = project
    cab.cabinet_type = ct
    project.cabinets = [cab]

    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert "PÓŁKI SZKLANE" in headings

    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    glass_table = None
    for idx, heading in enumerate(headings):
        if heading == "PÓŁKI SZKLANE":
            glass_table = body_tables[idx]
            break

    assert glass_table is not None
    assert len(glass_table.rows) == 2
    row = glass_table.rows[1].cells
    assert row[1].text == "półka szklana"
    assert row[3].text == "2"


def test_accessories_are_aggregated_without_sequence(tmp_path):
    """
    Given: two cabinets with the same accessory name
    When: generating AKCESORIA
    Then: accessory is summed into a single row and Poz. is a row index (not cabinet seq)
    """
    project = Project(
        name="Accessory Aggregate Project",
        kitchen_type="LOFT",
        order_number="ORM002",
        client_name="Client",
        client_address="Address",
        client_phone="555-000",
        client_email="client@example.com",
    )

    cabinet_1 = ProjectCabinet(
        sequence_number=1,
        body_color="Oak",
        front_color="Maple",
        handle_type="KROMA",
        quantity=2,
    )
    cabinet_2 = ProjectCabinet(
        sequence_number=2,
        body_color="Oak",
        front_color="Maple",
        handle_type="KROMA",
        quantity=1,
    )
    cabinet_1.project = project
    cabinet_2.project = project
    project.cabinets = [cabinet_1, cabinet_2]

    accessory_1 = Accessory(name="Hinge X")
    accessory_2 = Accessory(name="Hinge X")

    link_1 = ProjectCabinetAccessory(count=4)
    link_1.project_cabinet = cabinet_1
    link_1.accessory = accessory_1

    link_2 = ProjectCabinetAccessory(count=5)
    link_2.project_cabinet = cabinet_2
    link_2.accessory = accessory_2

    cabinet_1.accessories = [link_1]
    cabinet_2.accessories = [link_2]

    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)
    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]
    acc_table = body_tables[0]

    assert len(acc_table.rows) == 2
    cells = acc_table.rows[1].cells
    assert cells[0].text == "1"
    assert cells[1].text == "Hinge X"
    assert cells[2].text == str((4 * 2) + (5 * 1))
    assert cells[3].text == ""


def test_notes_and_footer(tmp_path, sample_project_orm):
    """
    Given: a generated report with notes
    When: reading paragraphs and footer
    Then: notes and page field exist
    """
    rg = ReportGenerator()
    output = rg.generate(sample_project_orm, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Blaty: Blaty test" in text
    assert "Cokoły: Cokoły test" in text
    assert "Uwagi: Uwagi test" in text
    footer = doc.sections[0].footer
    xml = footer.tables[0].cell(0, 1).paragraphs[0]._p.xml
    assert "fldSimple" in xml and "PAGE" in xml


@pytest.fixture
def project_with_custom_cabinets():
    """
    Given: a Project with both catalog and custom cabinets
    """
    # Project setup
    project = Project(
        name="Mixed Cabinet Project",
        kitchen_type="LOFT",
        order_number="MIX001",
        client_name="Mixed Client",
        client_address="789 Custom Ave",
        client_phone="555-MIX",
        client_email="mixed@example.com",
        blaty=False,
        cokoly=False,
        uwagi=False,
    )

    # Create catalog template for regular cabinet
    catalog_template = CabinetTemplate(
        kitchen_type="LOFT",
        name="D40",
    )

    # Add parts to catalog template
    catalog_parts = [
        CabinetPart(
            part_name="Catalog Panel",
            height_mm=600,
            width_mm=400,
            pieces=2,
            material="PLYTA 18",
        ),
        CabinetPart(
            part_name="Catalog Front",
            height_mm=598,
            width_mm=398,
            pieces=1,
            material="FRONT 18",
        ),
    ]
    catalog_template.parts = catalog_parts

    # Create catalog cabinet
    catalog_cabinet = ProjectCabinet(
        sequence_number=1,
        body_color="White",
        front_color="Oak",
        handle_type="Standard",
        quantity=1,
        cabinet_type=catalog_template,
    )

    # Create custom cabinet (type_id=NULL)
    custom_cabinet = ProjectCabinet(
        sequence_number=2,
        body_color="Gray",
        front_color="Black",
        handle_type="Push-to-open",
        quantity=2,
        cabinet_type=None,  # Custom cabinet
    )

    # Add custom parts
    custom_parts = [
        ProjectCabinetPart(
            part_name="Custom Side Panel",
            height_mm=720,
            width_mm=560,
            pieces=2,
            material="PLYTA 18",
            comments="Custom side panels",
        ),
        ProjectCabinetPart(
            part_name="Custom Front Door",
            height_mm=718,
            width_mm=558,
            pieces=1,
            material="FRONT 20",
            comments="Custom front door",
        ),
        ProjectCabinetPart(
            part_name="Custom Back Panel",
            height_mm=710,
            width_mm=550,
            pieces=1,
            material="HDF 3",
            comments="Custom HDF back",
        ),
    ]
    custom_cabinet.parts = custom_parts

    # Link cabinets to project
    project.cabinets = [catalog_cabinet, custom_cabinet]

    return project


def test_custom_cabinets_in_report(tmp_path, project_with_custom_cabinets):
    """
    Given: a project with both catalog and custom cabinets
    When: generating a report
    Then: both catalog and custom cabinet parts appear in appropriate sections
    """
    project = project_with_custom_cabinets
    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    # Get body tables (excluding header/footer)
    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    # With color grouping we expect multiple FORMATKI/FRONTY sections plus HDF.
    assert len(body_tables) >= 5

    # Check section headings
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any(h.startswith("FORMATKI (PLYTA 18) - ") for h in headings)
    assert any(h.startswith("FRONTY - ") for h in headings)
    assert "HDF" in headings
    assert "AKCESORIA" in headings

    seq1 = get_circled_number(1)
    seq2 = get_circled_number(2)

    def find_row_cells(part_name: str):
        for table in body_tables:
            for row in table.rows[1:]:
                if row.cells[1].text.strip() == part_name:
                    return row.cells
        return None

    catalog_panel = find_row_cells("Catalog Panel")
    custom_panel = find_row_cells("Custom Side Panel")
    assert catalog_panel is not None, "Catalog panel not found in FORMATKI"
    assert custom_panel is not None, "Custom side panel not found in FORMATKI"
    assert catalog_panel[0].text.strip() == seq1
    assert catalog_panel[3].text.strip() == "2"
    assert custom_panel[0].text.strip() == seq2
    assert custom_panel[3].text.strip() == "4"

    catalog_front = find_row_cells("Catalog Front")
    custom_front = find_row_cells("Custom Front Door")
    assert catalog_front is not None, "Catalog front not found in FRONTY"
    assert custom_front is not None, "Custom front door not found in FRONTY"
    assert catalog_front[0].text.strip() == seq1
    assert catalog_front[3].text.strip() == "1"
    assert catalog_front[5].text.strip() == "Oak"
    assert custom_front[0].text.strip() == seq2
    assert custom_front[3].text.strip() == "2"
    assert custom_front[5].text.strip() == "Black"

    custom_hdf = find_row_cells("Custom Back Panel")
    assert custom_hdf is not None, "Custom HDF back panel not found in HDF section"
    assert custom_hdf[0].text.strip() == seq2
    assert custom_hdf[3].text.strip() == "2"


def test_custom_cabinet_sequence_numbers(tmp_path, project_with_custom_cabinets):
    """
    Given: a project with custom cabinets having sequence numbers
    When: generating a report
    Then: custom cabinet parts have correct sequence number symbols
    """
    project = project_with_custom_cabinets
    rg = ReportGenerator()
    output = rg.generate(project, output_dir=str(tmp_path), auto_open=False)
    doc = Document(output)

    # Get body tables
    header_tables = doc.sections[0].header.tables
    footer_tables = doc.sections[0].footer.tables
    body_tables = [t for t in doc.tables if t not in header_tables + footer_tables]

    seq1 = get_circled_number(1)
    seq2 = get_circled_number(2)
    found_sequences = set()

    for table in body_tables:
        for row in table.rows[1:]:
            name = row.cells[1].text.strip()
            if name in ("Catalog Panel", "Custom Side Panel"):
                found_sequences.add(row.cells[0].text.strip())

    assert seq1 in found_sequences, (
        f"Sequence number {seq1} not found (catalog cabinet)"
    )
    assert seq2 in found_sequences, f"Sequence number {seq2} not found (custom cabinet)"
