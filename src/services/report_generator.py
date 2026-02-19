import os
import re
import sys
import subprocess
import logging
from pathlib import Path
from datetime import date
from typing import List, Tuple, Optional, Any, Dict
from types import SimpleNamespace

from docx import Document
from docx.document import Document as DocxDocument
from docx.section import Section
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from src.db_schema.orm_models import Project
from src.services.project_service import ProjectService
from sqlalchemy.orm import Session

# Configure logging
logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Generates a .docx report for a Cabplanner project.

    This class is responsible only for document generation and formatting.
    All data access and business logic is delegated to service classes.
    """

    def __init__(
        self,
        program_logo_path: Optional[Path] = None,
        company_logo_path: Optional[Path] = None,
        db_session: Optional[Session] = None,
    ) -> None:
        # Make sure paths are either None or proper Path objects
        self.program_logo_path = (
            program_logo_path
            if program_logo_path and not isinstance(program_logo_path, Session)
            else None
        )
        self.company_logo_path = (
            company_logo_path
            if company_logo_path and not isinstance(company_logo_path, Session)
            else None
        )
        self.db_session = db_session

        # Initialize ProjectService if session is provided
        self.project_service = ProjectService(db_session) if db_session else None

    def generate(
        self,
        project: Project,
        output_dir: str = "documents/reports",
        auto_open: bool = True,
    ) -> str:
        """
        Generate the .docx report and return its file path.
        """
        try:
            logger.info(
                f"Generating report for project: {project.name} (ID: {project.id})"
            )
            doc = Document()
            section = doc.sections[0]
            section.different_first_page_header_footer = False

            # Header and footer on every page
            self._add_header(section, project)
            self._add_footer(section)

            # Get data for report
            if self.project_service and project.id:
                # Get aggregated elements from service
                logger.debug("Using project service to get aggregated elements")
                elements = self.project_service.get_aggregated_project_elements(
                    project.id
                )
                formatki = self._dict_to_namespace_list(elements["formatki"])
                fronty = self._dict_to_namespace_list(elements["fronty"])
                witryny = self._dict_to_namespace_list(elements.get("witryny", []))
                hdf = self._dict_to_namespace_list(elements["hdf"])
                akcesoria = self._dict_to_namespace_list(elements["akcesoria"])
            else:
                # Fallback to direct data extraction if service not available
                logger.warning(
                    "Project service not available, using direct data extraction"
                )
                formatki, fronty, witryny, hdf, akcesoria = (
                    self._extract_elements_directly_with_witryny(project)
                )

            # Always sort by cabinet number first, then by color.
            formatki = self._sort_by_cabinet_and_color(formatki)
            fronty = self._sort_by_cabinet_and_color(fronty)
            witryny = self._sort_by_cabinet_and_color(witryny)
            hdf = self._sort_by_cabinet_and_color(hdf)
            akcesoria = self._aggregate_accessories(akcesoria)

            # Split formatki by material type
            formatki_plyta_12 = [
                p for p in formatki if getattr(p, "material", "") == "PLYTA 12"
            ]
            formatki_plyta_16 = [
                p for p in formatki if getattr(p, "material", "") == "PLYTA 16"
            ]
            formatki_plyta_18 = [
                p for p in formatki if getattr(p, "material", "") == "PLYTA 18"
            ]
            # Other formatki (legacy or without specific material)
            formatki_other = [
                p
                for p in formatki
                if getattr(p, "material", "")
                not in ("PLYTA 12", "PLYTA 16", "PLYTA 18")
            ]

            # Add sections - split FORMATKI by material type
            if formatki_plyta_18:
                self._add_parts_section(doc, "FORMATKI (PLYTA 18)", formatki_plyta_18)
            if formatki_plyta_16:
                self._add_parts_section(doc, "FORMATKI (PLYTA 16)", formatki_plyta_16)
            if formatki_plyta_12:
                self._add_parts_section(doc, "FORMATKI (PLYTA 12)", formatki_plyta_12)
            if formatki_other:
                self._add_parts_section(doc, "FORMATKI", formatki_other)
            self._add_parts_section(doc, "FRONTY", fronty)
            if witryny:
                self._add_parts_section(doc, "WITRYNY", witryny)
            self._add_parts_section(doc, "HDF", hdf)
            self._add_parts_section(doc, "AKCESORIA", akcesoria, accessory=True)

            # Optional notes
            self._add_notes(doc, project)

            # Save with handling for open files
            out_dir = Path(output_dir)
            out_dir.mkdir(parents=True, exist_ok=True)
            base_name = self._sanitize_filename_component(
                f"projekt_{project.order_number}"
            )
            output_path = self._get_available_filename(out_dir, base_name)
            doc.save(str(output_path))
            logger.info(f"Report saved to: {output_path}")

            if auto_open:
                self._open_file(output_path)

            return str(output_path)

        except Exception as e:
            logger.error(f"Error generating report: {str(e)}", exc_info=True)
            raise ReportGenerationError(f"Failed to generate report: {str(e)}")

    def _sort_by_cabinet_and_color(
        self, items: List[SimpleNamespace]
    ) -> List[SimpleNamespace]:
        """Sort items by cabinet sequence number and then by color."""
        return sorted(
            items,
            key=lambda item: (
                item.sequence,
                (getattr(item, "color", "") or "").strip().lower(),
                getattr(item, "name", "") or "",
            ),
        )

    def _dict_to_namespace_list(self, dict_list: List[Dict]) -> List[SimpleNamespace]:
        """Convert a list of dictionaries to a list of SimpleNamespace objects"""
        return [SimpleNamespace(**item) for item in dict_list]

    def _aggregate_accessories(
        self, accessories: List[SimpleNamespace]
    ) -> List[SimpleNamespace]:
        """
        Aggregate accessories across cabinets.
        Group by source_accessory_id when available, otherwise by normalized name + unit.
        """
        aggregated: Dict[tuple, SimpleNamespace] = {}

        for accessory in accessories:
            source_accessory_id = getattr(accessory, "source_accessory_id", None)
            name = (getattr(accessory, "name", "") or "").strip()
            unit = (getattr(accessory, "unit", "szt") or "szt").strip()
            notes = (getattr(accessory, "notes", "") or "").strip()
            quantity = int(getattr(accessory, "quantity", 0) or 0)

            if source_accessory_id is not None:
                key = ("id", source_accessory_id, unit)
            else:
                key = ("name", name.lower(), unit)

            if key not in aggregated:
                aggregated[key] = SimpleNamespace(
                    source_accessory_id=source_accessory_id,
                    name=name,
                    unit=unit,
                    quantity=0,
                    notes=notes,
                )

            aggregated[key].quantity += quantity

            if not aggregated[key].notes and notes:
                aggregated[key].notes = notes

        return sorted(
            aggregated.values(),
            key=lambda accessory: ((accessory.name or "").lower(), accessory.unit),
        )

    def _extract_elements_directly(
        self, project: Project
    ) -> Tuple[
        List[SimpleNamespace],
        List[SimpleNamespace],
        List[SimpleNamespace],
        List[SimpleNamespace],
    ]:
        """
        Backward-compatible extraction used by existing tests/callers.
        Returns: formatki, fronty, hdf, akcesoria.
        """
        formatki, fronty, witryny, hdf, akcesoria = (
            self._extract_elements_directly_with_witryny(project)
        )
        # Legacy API had no separate WITRYNY bucket; keep front-like semantics.
        return formatki, fronty + witryny, hdf, akcesoria

    def _extract_elements_directly_with_witryny(
        self, project: Project
    ) -> Tuple[
        List[SimpleNamespace],
        List[SimpleNamespace],
        List[SimpleNamespace],
        List[SimpleNamespace],
        List[SimpleNamespace],
    ]:
        """
        Extract parts lists directly from project object.
        This is a fallback method when ProjectService is not available.
        """
        formatki = []
        fronty = []
        witryny = []
        hdf = []
        akcesoria = []

        # This is a minimal implementation - ideally this logic should be in the service layer
        if not hasattr(project, "cabinets"):
            logger.warning("Project has no cabinets, returning empty lists")
            return formatki, fronty, witryny, hdf, akcesoria

        for cab in project.cabinets:
            qty = cab.quantity

            # Process both catalog and custom cabinets
            self._process_cabinet(cab, qty, formatki, fronty, hdf, witryny)

            # Process accessories for all cabinet types
            self._process_accessories(cab, qty, akcesoria)

        return formatki, fronty, witryny, hdf, akcesoria

    def _process_cabinet(self, cab, qty, formatki, fronty, hdf, witryny):
        """Process a cabinet (catalog or custom) and add its parts to the appropriate lists"""
        # Get the sequence number for this cabinet
        seq_num = getattr(cab, "sequence_number", 0)
        from src.services.project_service import get_circled_number

        seq_symbol = get_circled_number(seq_num)

        # Determine parts source based on cabinet type
        ct = cab.cabinet_type
        if ct:
            # Catalog cabinet: get parts from template
            parts = ct.parts
        else:
            # Custom cabinet: get parts directly from cabinet (snapshot architecture)
            parts = getattr(cab, "parts", [])

        # Process all parts
        for part in parts:
            part_qty = part.pieces * qty

            # Determine material
            material = part.material
            if not material:
                # For catalog cabinets, infer material from part name if not set
                if ct and not material:
                    part_name_lc = part.part_name.lower()
                    if "witryn" in part_name_lc:
                        material = "WITRYNA"
                    elif "front" in part_name_lc:
                        material = "FRONT"
                    elif "hdf" in part_name_lc:
                        material = "HDF"
                    else:
                        material = "PLYTA"  # Default for panels
                else:
                    # For custom cabinets, default to PLYTA
                    material = "PLYTA"

            # Determine category based on material (use startswith for variants like "FRONT 18", "HDF 3")
            if material and material.upper().startswith("WITRYNA"):
                witryny.append(
                    SimpleNamespace(
                        seq=seq_symbol,
                        sequence=seq_num,
                        name=part.part_name,
                        quantity=part_qty,
                        width=part.width_mm,
                        height=part.height_mm,
                        color=cab.front_color,
                        wrapping=getattr(part, "wrapping", "") or "",
                        notes=f"Handle: {cab.handle_type}",
                    )
                )
            elif material and material.upper().startswith("FRONT"):
                fronty.append(
                    SimpleNamespace(
                        seq=seq_symbol,
                        sequence=seq_num,
                        name=part.part_name,
                        quantity=part_qty,
                        width=part.width_mm,
                        height=part.height_mm,
                        color=cab.front_color,
                        wrapping=getattr(part, "wrapping", "") or "",
                        notes=f"Handle: {cab.handle_type}",
                    )
                )
            elif material and material.upper().startswith("HDF"):
                hdf.append(
                    SimpleNamespace(
                        seq=seq_symbol,
                        sequence=seq_num,
                        name=part.part_name,
                        quantity=part_qty,
                        width=part.width_mm,
                        height=part.height_mm,
                        color="",
                        wrapping=getattr(part, "wrapping", "") or "",
                        notes=part.comments or "",
                    )
                )
            else:
                # Default to formatki (panels)
                formatki.append(
                    SimpleNamespace(
                        seq=seq_symbol,
                        sequence=seq_num,
                        name=part.part_name,
                        quantity=part_qty,
                        width=part.width_mm,
                        height=part.height_mm,
                        color=cab.body_color,
                        material=material,
                        wrapping=getattr(part, "wrapping", "") or "",
                        notes=part.comments or "",
                    )
                )

    def _process_accessories(self, cab, qty, akcesoria):
        """Process accessories for a cabinet and add them to the accessories list"""
        for link in getattr(cab, "accessories", []):
            acc = link.accessory
            total = link.count * qty
            akcesoria.append(
                SimpleNamespace(
                    name=acc.name,
                    source_accessory_id=getattr(acc, "id", None),
                    unit=acc.unit,
                    quantity=total,
                    notes="",
                )
            )

    def _add_header(self, section: Section, project: Project) -> None:
        header = section.header
        usable_width = section.page_width - section.left_margin - section.right_margin
        table = header.add_table(1, 2, usable_width)
        table.autofit = True

        # Left: metadata in single line with Polish labels
        cell_meta = table.rows[0].cells[0]
        p = cell_meta.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Single line format with Polish labels
        info_text = (
            f"Klient: {project.client_name} | "
            f"Adres: {project.client_address} | "
            f"Tel.: {project.client_phone} | "
            f"Email: {project.client_email} | "
            f"Nr zamówienia: {project.order_number} | "
            f"Typ kuchni: {project.kitchen_type} | "
            f"Data raportu: {date.today().isoformat()}"
        )

        run = p.add_run(info_text)
        run.font.size = Pt(10)

        # Right: company logo
        cell_logo = table.rows[0].cells[1]
        if self.company_logo_path and os.path.exists(self.company_logo_path):
            run_logo = cell_logo.paragraphs[0].add_run()
            run_logo.add_picture(self.company_logo_path, width=Inches(1))

        # Program logo below
        if self.program_logo_path and os.path.exists(self.program_logo_path):
            p_logo = header.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_p = p_logo.add_run()
            run_p.add_picture(self.program_logo_path, width=Inches(0.75))

    def _add_footer(self, section: Section) -> None:
        footer = section.footer
        usable_width = section.page_width - section.left_margin - section.right_margin
        table = footer.add_table(1, 2, usable_width)
        table.autofit = True

        # Branding
        cell_brand = table.rows[0].cells[0]
        p_brand = cell_brand.paragraphs[0]
        p_brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_brand.add_run("Wygenerowano przez Cabplanner")
        run.italic = True

        # Page number
        cell_page = table.rows[0].cells[1]
        p_page = cell_page.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p_page._p.append(fld)

    def _add_parts_section(
        self, doc: DocxDocument, title: str, parts: List[Any], accessory: bool = False
    ) -> None:
        # Check if we need a page break before adding section
        if parts and self._should_break_page_for_section(doc, len(parts)):
            doc.add_page_break()

        doc.add_heading(title, level=2)
        if not parts:
            doc.add_paragraph("Brak pozycji.")
            return

        cols = (
            ["Poz.", "Nazwa akcesorium", "Ilość", "Jedn.", "Uwagi"]
            if accessory
            else [
                "Lp.",
                "Nazwa",
                "Wymiary (mm)",
                "Ilość",
                "Okleina",
                "Kolor",
                "Uwagi",
            ]
        )
        table = doc.add_table(rows=1, cols=len(cols))
        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = col
        qty_col_idx = 2 if accessory else 3
        hdr[qty_col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_index, part in enumerate(parts, start=1):
            cells = table.add_row().cells

            if accessory:
                # Accessories are project-wide aggregate; use row index, not cabinet sequence.
                cells[0].text = str(row_index)
                cells[1].text = part.name
                cells[2].text = str(part.quantity)
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[3].text = getattr(part, "unit", "szt") or "szt"
                cells[4].text = getattr(part, "notes", "") or ""
            else:
                # Parts keep cabinet sequence marker.
                cells[0].text = getattr(part, "seq", "")
                cells[1].text = part.name
                cells[2].text = f"{part.width} x {part.height}"
                cells[3].text = str(part.quantity)
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[4].text = getattr(part, "wrapping", "") or ""
                cells[5].text = getattr(part, "color", "") or ""
                cells[6].text = getattr(part, "notes", "") or ""

    def _add_notes(self, doc: DocxDocument, project: Project) -> None:
        if getattr(project, "blaty_note", None):
            p = doc.add_paragraph()
            p.add_run("Blaty: ").bold = True
            p.add_run(project.blaty_note)
        if getattr(project, "cokoly_note", None):
            p = doc.add_paragraph()
            p.add_run("Cokoły: ").bold = True
            p.add_run(project.cokoly_note)
        if getattr(project, "uwagi_note", None):
            p = doc.add_paragraph()
            p.add_run("Uwagi: ").bold = True
            p.add_run(project.uwagi_note)

    def _get_available_filename(self, output_dir: Path, base_name: str) -> Path:
        """
        Return the first free filename.
        Examples: raport.docx, raport (1).docx, raport (2).docx.
        """
        safe_base_name = self._sanitize_filename_component(base_name)

        for i in range(0, 100):
            suffix = "" if i == 0 else f" ({i})"
            candidate = output_dir / f"{safe_base_name}{suffix}.docx"
            if not candidate.exists():
                if i > 0:
                    logger.info(
                        f"Report file already exists, using alternative name: {candidate.name}"
                    )
                return candidate

        raise OSError(
            f"Nie można utworzyć unikalnej nazwy raportu dla '{safe_base_name}.docx'."
        )

    def _sanitize_filename_component(self, name: str) -> str:
        """Normalize file name to be safe on common file systems."""
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", (name or "").strip())
        sanitized = re.sub(r"\s+", " ", sanitized).strip(" .")
        return sanitized or "raport"

    def _should_break_page_for_section(
        self, doc: DocxDocument, items_count: int
    ) -> bool:
        """Check if we should add a page break before a new section."""
        if items_count == 0:
            return False

        # Rough calculation: if we have less than 6 lines remaining on page
        # (header + 2 table header rows + at least 1 data row + some margin)
        # This is a simple heuristic - Word's actual pagination is more complex
        current_elements = len(doc.paragraphs) + sum(
            len(table.rows) for table in doc.tables
        )

        # Assume ~50 lines per page, break if we have less than 6 lines left
        lines_used = current_elements % 50
        lines_remaining = 50 - lines_used

        # Need at least: title (1) + table header (1) + 1 data row (1) + margin (3) = 6 lines
        min_lines_needed = 6

        return lines_remaining < min_lines_needed

    def _open_file(self, path: Path) -> None:
        try:
            if os.name == "nt":
                os.startfile(path)
            elif sys.platform == "darwin":
                subprocess.call(["open", path])
            else:
                subprocess.call(["xdg-open", path])
            logger.debug(f"Successfully opened file: {path}")
        except Exception as e:
            logger.error(f"Error opening file {path}: {str(e)}")
            # Don't raise - this is a non-critical operation


class ReportGenerationError(Exception):
    """Exception raised when report generation fails"""

    pass
